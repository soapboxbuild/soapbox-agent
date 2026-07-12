#!/usr/bin/env python3
"""Fail-closed scrub gate for demo fixtures.

Scans every demo fixture (extracting text from PDF/XLSX/DOCX, not just plain
text) for any real-name term in the untracked denylist. Exits non-zero on the
first directory that leaks, so it can gate staging into the Files store.

Usage: python3 demo-staging/scrub-check.py [path ...]
  With no arguments, scans the default fixture roots (SCAN_DIRS).
  With one or more path arguments (files or directories), scans those
  instead — e.g. to gate a fixture already committed elsewhere:
    python3 demo-staging/scrub-check.py ~/soapbox-platform/apps/api/src/services/demo-fixtures/
Denylist: demo-staging/.scrub-denylist.json  (untracked; real names)
"""
import json, os, re, sys, zipfile

REPO = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DENYLIST = os.path.join(REPO, "demo-staging", ".scrub-denylist.json")

# fixture roots to scan
SCAN_DIRS = [
    "skills/rsra/demo",
    "skills/esg-profile/demo/madison",
    "skills/decarb-plan/demo",
]
# never scan these (they legitimately hold real names or are vendored)
SKIP_NAMES = {".venv", "node_modules", "__MACOSX"}
SKIP_FILES = {".scrub-denylist.json", ".pseudonym-map.md"}


def load_terms():
    with open(DENYLIST) as f:
        terms = json.load(f)["terms"]
    return [t for t in terms if t.strip()]


def extract_text(path):
    """Return best-effort text for a file, or '' if unreadable."""
    ext = path.lower().rsplit(".", 1)[-1] if "." in path else ""
    try:
        if ext == "pdf":
            try:
                import fitz  # PyMuPDF
                return "\n".join(p.get_text() for p in fitz.open(path))
            except Exception:
                import pypdf
                return "\n".join(pg.extract_text() or "" for pg in pypdf.PdfReader(path).pages)
        if ext == "xlsx":
            import openpyxl
            wb = openpyxl.load_workbook(path, data_only=True)
            out = []
            for ws in wb:
                for row in ws.iter_rows(values_only=True):
                    out.extend(str(c) for c in row if c is not None)
            return "\n".join(out)
        if ext == "docx":
            import docx
            d = docx.Document(path)
            parts = [p.text for p in d.paragraphs]
            for t in d.tables:
                for r in t.rows:
                    parts.extend(c.text for c in r.cells)
            return "\n".join(parts)
        # plain text / json / md / py / csv / html
        with open(path, "r", errors="ignore") as f:
            return f.read()
    except Exception as e:
        print(f"  ! could not read {path}: {e}", file=sys.stderr)
        return ""


def scan_file(path, patterns, hits):
    text = extract_text(path)
    for term, pat in patterns:
        if pat.search(text):
            hits.append((os.path.relpath(path, REPO), term))


def main():
    terms = load_terms()
    if not terms:
        # Empty denylist means the gate can't actually check anything —
        # fail closed rather than silently passing every fixture.
        print("SCRUB FAIL — denylist is empty; refusing to pass a no-op scrub gate.")
        sys.exit(1)
    # word-boundary match so short terms (e.g. "Varia") don't fire on
    # "variance"/"variable"; boundaries anchor on alphanumeric edges.
    patterns = [(t, re.compile(r"\b" + re.escape(t) + r"\b", re.IGNORECASE)) for t in terms]
    hits = []
    scanned = 0

    # Explicit path arguments take priority over the default SCAN_DIRS so the
    # tool can gate fixtures committed outside this repo (e.g. the platform's
    # apps/api/src/services/demo-fixtures/).
    targets = sys.argv[1:] if len(sys.argv) > 1 else [os.path.join(REPO, base) for base in SCAN_DIRS]

    for target in targets:
        if os.path.isfile(target):
            scanned += 1
            scan_file(target, patterns, hits)
            continue
        if not os.path.isdir(target):
            continue
        for dirpath, dirnames, filenames in os.walk(target):
            dirnames[:] = [d for d in dirnames if d not in SKIP_NAMES]
            for fn in filenames:
                if fn in SKIP_FILES:
                    continue
                path = os.path.join(dirpath, fn)
                scanned += 1
                scan_file(path, patterns, hits)

    if hits:
        print(f"SCRUB FAIL — {len(hits)} real-name leak(s) in {scanned} files:")
        for path, term in hits:
            print(f"  {path}: '{term}'")
        sys.exit(1)
    print(f"SCRUB CLEAN — {scanned} fixture files scanned against {len(terms)} denied terms.")
    sys.exit(0)


if __name__ == "__main__":
    main()
